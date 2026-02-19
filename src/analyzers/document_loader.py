"""
Document Loader — Format-agnostic document ingestion.

Supports: .txt, .md, .json, .html, .pdf, .docx
Produces: List of DocumentChunk objects with text content, source metadata,
          and position info for downstream analysis.
"""

from __future__ import annotations

import json
from dataclasses import dataclass, field
from pathlib import Path

from src.core.logging import get_logger

logger = get_logger(__name__)


# ============================================================
# Models
# ============================================================

@dataclass
class DocumentChunk:
    """A chunk of text from a loaded document."""
    text: str
    source_file: str
    chunk_index: int = 0
    total_chunks: int = 1
    file_type: str = ""
    metadata: dict = field(default_factory=dict)

    @property
    def char_count(self) -> int:
        return len(self.text)

    @property
    def word_count(self) -> int:
        return len(self.text.split())


@dataclass
class LoadedDocument:
    """A fully loaded document with all its chunks."""
    source_file: str
    file_type: str
    chunks: list[DocumentChunk] = field(default_factory=list)
    load_error: str | None = None
    metadata: dict = field(default_factory=dict)

    @property
    def total_text(self) -> str:
        return "\n\n".join(c.text for c in self.chunks)

    @property
    def total_chars(self) -> int:
        return sum(c.char_count for c in self.chunks)

    @property
    def total_words(self) -> int:
        return sum(c.word_count for c in self.chunks)

    @property
    def loaded_successfully(self) -> bool:
        return self.load_error is None and len(self.chunks) > 0


@dataclass
class LoadResult:
    """Result of loading one or more documents."""
    documents: list[LoadedDocument] = field(default_factory=list)
    total_files: int = 0
    successful: int = 0
    failed: int = 0
    errors: list[str] = field(default_factory=list)

    @property
    def all_text(self) -> str:
        return "\n\n---\n\n".join(
            d.total_text for d in self.documents if d.loaded_successfully
        )

    @property
    def all_chunks(self) -> list[DocumentChunk]:
        chunks = []
        for d in self.documents:
            chunks.extend(d.chunks)
        return chunks


# ============================================================
# Supported file types
# ============================================================

SUPPORTED_EXTENSIONS = {".txt", ".md", ".json", ".html", ".htm", ".pdf", ".docx"}


# ============================================================
# Document Loader
# ============================================================

class DocumentLoader:
    """
    Format-agnostic document loader.

    Loads files from a directory or individual paths, extracts text content,
    and chunks large documents for downstream LLM analysis.
    """

    def __init__(
        self,
        chunk_size: int = 4000,
        chunk_overlap: int = 200,
        max_file_size_mb: float = 50.0,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.max_file_size_bytes = int(max_file_size_mb * 1024 * 1024)

    def load_directory(self, dir_path: str | Path, recursive: bool = True) -> LoadResult:
        """Load all supported documents from a directory."""
        dp = Path(dir_path)
        if not dp.exists():
            return LoadResult(errors=[f"Directory not found: {dir_path}"])
        if not dp.is_dir():
            return LoadResult(errors=[f"Not a directory: {dir_path}"])

        files = []
        if recursive:
            for ext in SUPPORTED_EXTENSIONS:
                files.extend(dp.rglob(f"*{ext}"))
        else:
            for ext in SUPPORTED_EXTENSIONS:
                files.extend(dp.glob(f"*{ext}"))

        # Sort for deterministic order
        files = sorted(set(files))

        if not files:
            return LoadResult(errors=[f"No supported files found in {dir_path}"])

        return self.load_files(files)

    def load_files(self, file_paths: list[str | Path]) -> LoadResult:
        """Load a list of specific files."""
        result = LoadResult(total_files=len(file_paths))

        for fp in file_paths:
            doc = self.load_file(fp)
            result.documents.append(doc)
            if doc.loaded_successfully:
                result.successful += 1
            else:
                result.failed += 1
                if doc.load_error:
                    result.errors.append(f"{fp}: {doc.load_error}")

        logger.info(
            "documents_loaded",
            total=result.total_files,
            successful=result.successful,
            failed=result.failed,
        )

        return result

    def load_file(self, file_path: str | Path) -> LoadedDocument:
        """Load a single file, auto-detecting format."""
        fp = Path(file_path)

        if not fp.exists():
            return LoadedDocument(
                source_file=str(fp),
                file_type="unknown",
                load_error=f"File not found: {fp}",
            )

        if fp.suffix.lower() not in SUPPORTED_EXTENSIONS:
            return LoadedDocument(
                source_file=str(fp),
                file_type=fp.suffix,
                load_error=f"Unsupported file type: {fp.suffix}",
            )

        # Check file size
        if fp.stat().st_size > self.max_file_size_bytes:
            return LoadedDocument(
                source_file=str(fp),
                file_type=fp.suffix,
                load_error=f"File too large: {fp.stat().st_size / 1024 / 1024:.1f}MB (max {self.max_file_size_bytes / 1024 / 1024:.0f}MB)",
            )

        try:
            text = self._extract_text(fp)
            if not text or not text.strip():
                return LoadedDocument(
                    source_file=str(fp),
                    file_type=fp.suffix,
                    load_error="File is empty or contains no extractable text",
                )

            chunks = self._chunk_text(text, str(fp), fp.suffix)

            return LoadedDocument(
                source_file=str(fp),
                file_type=fp.suffix,
                chunks=chunks,
                metadata={"size_bytes": fp.stat().st_size, "name": fp.name},
            )

        except Exception as e:
            logger.error("document_load_error", file=str(fp), error=str(e))
            return LoadedDocument(
                source_file=str(fp),
                file_type=fp.suffix,
                load_error=str(e),
            )

    # ============================================================
    # Format-specific extractors
    # ============================================================

    def _extract_text(self, fp: Path) -> str:
        """Extract text from file based on extension."""
        ext = fp.suffix.lower()

        if ext in (".txt", ".md"):
            return self._load_text(fp)
        elif ext == ".json":
            return self._load_json(fp)
        elif ext in (".html", ".htm"):
            return self._load_html(fp)
        elif ext == ".pdf":
            return self._load_pdf(fp)
        elif ext == ".docx":
            return self._load_docx(fp)
        else:
            raise ValueError(f"No extractor for: {ext}")

    def _load_text(self, fp: Path) -> str:
        """Load plain text / markdown."""
        return fp.read_text(encoding="utf-8", errors="replace")

    def _load_json(self, fp: Path) -> str:
        """Load JSON, converting to readable text."""
        with open(fp, encoding="utf-8") as f:
            data = json.load(f)

        if isinstance(data, str):
            return data
        elif isinstance(data, list):
            parts = []
            for item in data:
                if isinstance(item, str):
                    parts.append(item)
                elif isinstance(item, dict):
                    parts.append(self._dict_to_text(item))
                else:
                    parts.append(str(item))
            return "\n\n".join(parts)
        elif isinstance(data, dict):
            return self._dict_to_text(data)
        else:
            return str(data)

    def _dict_to_text(self, d: dict, indent: int = 0) -> str:
        """Convert a dict to human-readable text."""
        lines = []
        prefix = "  " * indent
        for key, value in d.items():
            if isinstance(value, dict):
                lines.append(f"{prefix}{key}:")
                lines.append(self._dict_to_text(value, indent + 1))
            elif isinstance(value, list):
                lines.append(f"{prefix}{key}:")
                for item in value:
                    if isinstance(item, dict):
                        lines.append(self._dict_to_text(item, indent + 1))
                    else:
                        lines.append(f"{prefix}  - {item}")
            else:
                lines.append(f"{prefix}{key}: {value}")
        return "\n".join(lines)

    def _load_html(self, fp: Path) -> str:
        """Load HTML, stripping tags to extract text."""
        raw = fp.read_text(encoding="utf-8", errors="replace")
        # Simple tag stripping (no BeautifulSoup dependency)
        import re
        # Remove script and style blocks
        text = re.sub(r"<script[^>]*>.*?</script>", "", raw, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
        # Remove HTML tags
        text = re.sub(r"<[^>]+>", " ", text)
        # Decode common entities
        text = text.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
        text = text.replace("&nbsp;", " ").replace("&quot;", '"')
        # Clean up whitespace
        text = re.sub(r"\s+", " ", text).strip()
        return text

    def _load_pdf(self, fp: Path) -> str:
        """Load PDF using PyPDF2."""
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError(
                "PyPDF2 is required to load PDF files. "
                "Install with: pip install PyPDF2"
            )

        reader = PdfReader(str(fp))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                pages.append(text.strip())

        return "\n\n".join(pages)

    def _load_docx(self, fp: Path) -> str:
        """Load DOCX using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required to load DOCX files. "
                "Install with: pip install python-docx"
            )

        doc = Document(str(fp))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n\n".join(paragraphs)

    # ============================================================
    # Chunking
    # ============================================================

    def _chunk_text(
        self,
        text: str,
        source_file: str,
        file_type: str,
    ) -> list[DocumentChunk]:
        """Split text into overlapping chunks for LLM processing."""
        if len(text) <= self.chunk_size:
            return [DocumentChunk(
                text=text,
                source_file=source_file,
                chunk_index=0,
                total_chunks=1,
                file_type=file_type,
            )]

        chunks = []
        start = 0
        chunk_index = 0

        while start < len(text):
            end = start + self.chunk_size

            # Try to break at a paragraph or sentence boundary
            if end < len(text):
                # Look for paragraph break
                para_break = text.rfind("\n\n", start, end)
                if para_break > start + self.chunk_size // 2:
                    end = para_break + 2
                else:
                    # Look for sentence break
                    for sep in (". ", ".\n", "! ", "? "):
                        sent_break = text.rfind(sep, start, end)
                        if sent_break > start + self.chunk_size // 2:
                            end = sent_break + len(sep)
                            break

            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append(DocumentChunk(
                    text=chunk_text,
                    source_file=source_file,
                    chunk_index=chunk_index,
                    file_type=file_type,
                ))
                chunk_index += 1

            start = end - self.chunk_overlap
            if start >= len(text):
                break

        # Update total_chunks
        for c in chunks:
            c.total_chunks = len(chunks)

        return chunks